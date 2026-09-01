"""Document parsing — extract UTF-8 text from uploaded bytes, per file type.

This is the CONTROL-PLANE (torch-free) text-extraction stage between raw upload and
onboarding. The onboard/train path (routers/jobs.py) must hand the ML plane extracted
TEXT, never raw PDF/DOCX bytes — so every uploaded document runs through extract_text()
and the result is what gets persisted + read back for retrieval and onboarding.

Pure-pip and torch-free: pypdf / python-docx / beautifulsoup4 for the core formats, plus
openpyxl+xlrd (spreadsheets), striprtf (RTF-flavoured .doc) and — only as a LAST resort for
scanned PDFs — RapidOCR (ONNX, CPU) rendered via pypdfium2. No torch, no ML deps, no
system-binary HARD dependency (LibreOffice for genuine legacy .doc is optional; absent, that
one format fails cleanly asking for a .docx). A scanned PDF now OCRs at upload instead of
silently yielding nothing.

Contract: extract_text(filename, data) -> (text, ok, error).
  ok=True  -> `text` is the extracted UTF-8 body ("" is allowed only for genuinely empty
              text files; an image-only PDF that yields nothing is ok=False).
  ok=False -> `error` is a SHORT human-readable reason (unsupported type, encrypted,
              corrupt, decode failure). `text` is "" and the caller marks the doc failed.
"""
from __future__ import annotations

import csv
import io
import logging
import os
import shutil
import subprocess
import tempfile
import zipfile

logger = logging.getLogger(__name__)

# Decompression-bomb guards (anti-DoS): a small DOCX (a zip) can inflate to gigabytes and a PDF can
# declare an enormous page count, either of which would exhaust memory/CPU during extraction. Refuse
# up front as a clean parse FAILURE (never an exception the caller has to catch). Generous ceilings —
# a legitimate document is orders of magnitude under these.
_MAX_DOCX_UNCOMPRESSED_BYTES = 200 * 1024 * 1024  # 200 MB of inflated zip members
_MAX_DOCX_MEMBERS = 2_000                          # a real .docx has tens of parts, not thousands
_MAX_PDF_PAGES = 2_000

# OCR fallback (scanned PDFs). We only OCR when the text layer is (near-)empty — a real text PDF
# never touches this path. Render at ~200 DPI (RENDER_SCALE = 200/72) which is legible for RapidOCR
# without ballooning the bitmap. Below OCR_MIN_TEXT_CHARS of extracted text we treat the PDF as
# image-only and fall back to render+OCR.
_OCR_MIN_TEXT_CHARS = 20
_OCR_RENDER_SCALE = 200 / 72  # pypdfium2 scale == render DPI / 72 (its base is 72 DPI)

# Spreadsheet shape caps (retrieval/LLM-friendly, anti-DoS): a huge sheet is TRUNCATED, not failed.
_XLSX_MAX_ROWS_PER_SHEET = 5_000
_SHEET_MAX_TOTAL_CHARS = 500_000

# OLE (legacy binary) .doc magic and the LibreOffice headless-convert timeout.
_OLE_MAGIC = b"\xd0\xcf\x11\xe0"
_SOFFICE_TIMEOUT_S = 60

# Extensions we can extract text from. The upload validation (routers/corpora.py) and this
# module share this set as the single source of truth for "supported document type".
SUPPORTED_EXTS: frozenset[str] = frozenset(
    {".txt", ".md", ".pdf", ".docx", ".doc", ".html", ".htm",
     ".xlsx", ".xls", ".csv", ".tsv"}
)


class _ParseRejected(Exception):
    """A document we deliberately refuse to parse (e.g. a decompression bomb). The message is a short,
    safe reason; extract_text's normalizer turns it into (ok=False, error) like any other failure."""


def _ext(filename: str) -> str:
    """Lowercased extension including the dot ('' when the name has none)."""
    name = filename.rsplit("/", 1)[-1]
    return ("." + name.rsplit(".", 1)[-1].lower()) if "." in name else ""


def _extract_txt(data: bytes) -> str:
    """Plain text / markdown: decode as UTF-8. Fall back through common encodings so a
    Windows-authored (cp1252) or Latin-1 file doesn't fail the whole upload; last resort
    is a lossy UTF-8 decode (never raises) so we always return SOMETHING for a text file."""
    for enc in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


# RapidOCR engine is a heavy singleton: constructing it loads several ONNX models (seconds) and a
# module-level import would slow every process boot even for tenants who never upload a scan. So we
# import lazily and cache the one engine on first OCR use. `False` means "tried and failed to init"
# (e.g. missing native wheel) — we don't retry per document, we just skip OCR gracefully.
_ocr_engine = None  # None = not yet tried; False = init failed; else the RapidOCR instance


def _get_ocr_engine():
    """Lazily import + construct the RapidOCR engine ONCE (cached). Returns the engine, or None if
    it can't initialize in this environment (so the caller degrades to 'no extractable text' instead
    of crashing)."""
    global _ocr_engine
    if _ocr_engine is not None:
        return _ocr_engine or None
    try:
        from rapidocr_onnxruntime import RapidOCR
        _ocr_engine = RapidOCR()
    except Exception as exc:  # noqa: BLE001 — missing wheel / init error: OCR simply unavailable here
        logger.warning("RapidOCR unavailable, scanned-PDF OCR disabled: %s", exc)
        _ocr_engine = False
        return None
    return _ocr_engine


def _ocr_pdf(data: bytes) -> str:
    """Render each PDF page to a bitmap (pypdfium2, ~200 DPI) and OCR it with RapidOCR, joining the
    per-page text. Only called when the text layer was (near-)empty. Never raises: any engine error
    degrades to '' so the caller reports a clean parse failure, not a 500.

    Enforces OCR_MAX_PAGES BEFORE doing any OCR work — OCR runs inside the upload request at
    ~1-2s/page, so a big scan is refused (as a _ParseRejected) rather than blocking for minutes."""
    from . import config

    import pypdfium2 as pdfium

    engine = _get_ocr_engine()
    if engine is None:
        return ""

    pdf = pdfium.PdfDocument(data)
    try:
        n_pages = len(pdf)
        if n_pages > config.OCR_MAX_PAGES:
            raise _ParseRejected(
                f"scanned PDF has {n_pages} pages; OCR at upload is limited to "
                f"{config.OCR_MAX_PAGES} — split the file or upload a text PDF")
        pages_text: list[str] = []
        for i in range(n_pages):
            page = pdf[i]
            try:
                bitmap = page.render(scale=_OCR_RENDER_SCALE)
                img = bitmap.to_pil()
            finally:
                page.close()
            try:
                # RapidOCR takes a numpy array; PIL image -> np via the engine's own path. Result is
                # (list[[box, text, score]], elapse) or (None, ...) when nothing is detected.
                import numpy as np
                result, _ = engine(np.asarray(img))
            except Exception as exc:  # noqa: BLE001 — one page's OCR failing shouldn't lose the rest
                logger.info("OCR failed on page %d: %s", i, exc)
                continue
            if result:
                pages_text.append("\n".join(line[1] for line in result if line and len(line) > 1))
        return "\n\n".join(t for t in pages_text if t.strip())
    finally:
        pdf.close()


def _extract_pdf(data: bytes) -> str:
    """PDF via pypdf. Raises on an encrypted PDF we can't open with the empty password, or
    on a corrupt file — the caller turns that into ok=False + a short error. When the text layer is
    (near-)empty (a scanned/image-only PDF) we fall back to rendering + OCR (see _ocr_pdf); only then
    is empty output a genuine failure."""
    from pypdf import PdfReader
    from pypdf.errors import PdfReadError

    try:
        reader = PdfReader(io.BytesIO(data))
        if reader.is_encrypted:
            # Try the empty password (many "encrypted" PDFs are just permission-flagged);
            # a real password-protected file raises and becomes a clean parse failure.
            try:
                reader.decrypt("")
            except Exception as exc:  # noqa: BLE001 — normalize to our error contract
                raise PdfReadError("encrypted PDF (password required)") from exc
        # Refuse an absurd page count BEFORE iterating pages (each extract_text() is real work).
        n_pages = len(reader.pages)
        if n_pages > _MAX_PDF_PAGES:
            raise _ParseRejected(f"PDF has {n_pages} pages (limit {_MAX_PDF_PAGES})")
        text = "\n\n".join(page.extract_text() or "" for page in reader.pages)
    except _ParseRejected:
        raise
    except PdfReadError:
        raise
    except Exception as exc:  # noqa: BLE001 — pypdf raises assorted low-level errors on corrupt input
        raise PdfReadError(f"unreadable PDF: {type(exc).__name__}") from exc

    # (Near-)empty text layer -> this is a scan; render + OCR. _ocr_pdf may raise _ParseRejected
    # (page cap) which propagates as a clean failure; any other OCR error inside it returns "".
    if len(text.strip()) < _OCR_MIN_TEXT_CHARS:
        ocr_text = _ocr_pdf(data)
        if ocr_text.strip():
            return ocr_text
    return text


def _guard_zip(data: bytes, kind: str) -> None:
    """Decompression-bomb guard shared by every zip-backed format (DOCX and XLSX are both zips):
    inspect the central directory ONLY (no inflation) and raise _ParseRejected on too many members
    or a declared uncompressed total over the ceiling. `kind` names the format for the error."""
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            infos = zf.infolist()
            if len(infos) > _MAX_DOCX_MEMBERS:
                raise _ParseRejected(f"{kind} has {len(infos)} members (limit {_MAX_DOCX_MEMBERS})")
            total = sum(i.file_size for i in infos)  # declared uncompressed sizes
            if total > _MAX_DOCX_UNCOMPRESSED_BYTES:
                raise _ParseRejected(
                    f"{kind} inflates to {total} bytes (limit {_MAX_DOCX_UNCOMPRESSED_BYTES})")
    except zipfile.BadZipFile as exc:
        # Not a valid zip (e.g. a legacy .doc renamed) — a clean parse failure, not a crash.
        raise _ParseRejected(f"not a valid {kind} (bad zip)") from exc


def _extract_docx(data: bytes) -> str:
    """DOCX via python-docx. Pulls paragraph text plus table cell text (so a doc that puts
    its content in tables isn't onboarded empty). A non-DOCX (e.g. legacy .doc) or corrupt
    file raises PackageNotFoundError, which the caller turns into a clean parse failure."""
    from docx import Document as DocxDocument

    # A DOCX is a zip — refuse a decompression bomb (member count / declared uncompressed size)
    # BEFORE python-docx unzips anything.
    _guard_zip(data, "DOCX")

    doc = DocxDocument(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def _extract_html(data: bytes) -> str:
    """HTML via BeautifulSoup: strip tags, drop script/style, keep visible text. Uses the
    stdlib html.parser (no lxml dependency) so it stays pure-python."""
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(_extract_txt(data), "html.parser")
    for tag in soup(["script", "style"]):
        tag.decompose()
    return soup.get_text(separator="\n")


# --- spreadsheets ---------------------------------------------------------------------------
# Shape all spreadsheet formats into the SAME retrieval/LLM-friendly text: per sheet a
# "## Sheet: <name>" header, then one line per non-empty row with cells joined by " | ". A single
# writer (_rows_to_text) enforces the row/char caps and the "[truncated: N more rows]" marker so
# xlsx/xls/csv all truncate identically (truncation is fine; failing a big sheet is not).

def _fmt_cell(v: object) -> str:
    """One cell -> a trimmed string. None (empty cell) becomes '' so empty trailing cells vanish."""
    return "" if v is None else str(v).strip()


def _rows_to_text(sheets: list[tuple[str, "iter"]]) -> str:  # noqa: F821 — iterables of row-lists
    """Render (sheet_name, rows) pairs to text. Skips fully-empty rows; caps at
    _XLSX_MAX_ROWS_PER_SHEET per sheet and _SHEET_MAX_TOTAL_CHARS total, appending a
    '[truncated: N more rows]' marker past a per-sheet cap and stopping cleanly at the char cap."""
    out: list[str] = []
    total_chars = 0
    for name, rows in sheets:
        header = f"## Sheet: {name}"
        out.append(header)
        total_chars += len(header) + 1
        emitted = 0
        truncated_extra = 0
        for row in rows:
            cells = [_fmt_cell(c) for c in row]
            if not any(cells):
                continue  # fully-empty row: skip
            if emitted >= _XLSX_MAX_ROWS_PER_SHEET:
                truncated_extra += 1
                continue
            line = " | ".join(cells).rstrip(" |")
            if total_chars + len(line) + 1 > _SHEET_MAX_TOTAL_CHARS:
                out.append(f"[truncated: char limit {_SHEET_MAX_TOTAL_CHARS} reached]")
                return "\n".join(out)
            out.append(line)
            total_chars += len(line) + 1
            emitted += 1
        if truncated_extra:
            out.append(f"[truncated: {truncated_extra} more rows]")
    return "\n".join(out)


def _extract_xlsx(data: bytes) -> str:
    """XLSX via openpyxl (read_only, values only). A .xlsx is a zip — run the bomb guard first."""
    import openpyxl

    _guard_zip(data, "XLSX")
    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    try:
        sheets = [(ws.title, ws.iter_rows(values_only=True)) for ws in wb.worksheets]
        return _rows_to_text(sheets)
    finally:
        wb.close()


def _extract_xls(data: bytes) -> str:
    """Legacy .xls via xlrd (2.x is .xls-only by design — exactly its purpose)."""
    import xlrd

    book = xlrd.open_workbook(file_contents=data)
    sheets = []
    for sh in book.sheets():
        rows = (sh.row_values(r) for r in range(sh.nrows))
        sheets.append((sh.name, rows))
    return _rows_to_text(sheets)


def _extract_csv(data: bytes, ext: str) -> str:
    """CSV/TSV via the stdlib csv module. Decode with the same multi-encoding fallback the .txt path
    uses, then pick the delimiter: tab for .tsv, else sniff (comma/semicolon/tab), defaulting to
    comma. One sheet named for the extension, run through the shared row renderer."""
    text = _extract_txt(data)
    if ext == ".tsv":
        delimiter = "\t"
    else:
        try:
            delimiter = csv.Sniffer().sniff(text[:4096], delimiters=",;\t").delimiter
        except csv.Error:
            delimiter = ","
    reader = csv.reader(io.StringIO(text), delimiter=delimiter)
    label = "TSV" if ext == ".tsv" else "CSV"
    return _rows_to_text([(label, reader)])


# --- legacy .doc ----------------------------------------------------------------------------

def _soffice_to_docx(data: bytes) -> bytes | None:
    """Convert a genuine OLE binary .doc to .docx with headless LibreOffice, IF `soffice` is on PATH.
    Returns the .docx bytes, or None when LibreOffice isn't installed (dev boxes) or the convert
    fails. We never hand-parse the OLE format — LibreOffice is the only supported path.

    DEPLOY-CONTAINER TODO: bake `libreoffice`/`soffice` into the runtime image so genuine .doc files
    convert in production. Absent it (e.g. a Windows dev box) this returns None and the caller emits
    the 'save it as .docx' failure — that's the intended graceful degradation, not a bug."""
    soffice = shutil.which("soffice") or shutil.which("soffice.exe")
    if not soffice:
        return None
    with tempfile.TemporaryDirectory() as tmp:
        src = os.path.join(tmp, "in.doc")
        with open(src, "wb") as fh:
            fh.write(data)
        try:
            subprocess.run(
                [soffice, "--headless", "--convert-to", "docx", "--outdir", tmp, src],
                capture_output=True, timeout=_SOFFICE_TIMEOUT_S, check=True,
            )
        except Exception as exc:  # noqa: BLE001 — timeout / nonzero exit / missing profile: give up
            logger.info("soffice .doc->docx conversion failed: %s", exc)
            return None
        out = os.path.join(tmp, "in.docx")
        if not os.path.exists(out):
            return None
        with open(out, "rb") as fh:
            return fh.read()


def _extract_doc(data: bytes) -> str:
    """Legacy .doc — but many real-world ".doc" files are mislabeled, so SNIFF the bytes first:
      (a) 'PK'      -> actually a docx (zip) -> parse as docx
      (b) '{\\rtf'  -> RTF -> striprtf
      (c) OLE magic -> genuine binary .doc -> LibreOffice headless convert (if available)
      (d) otherwise -> clear failure asking for a .docx re-save.
    Never attempts to hand-parse the OLE binary format."""
    head = data[:8]
    if head[:2] == b"PK":
        return _extract_docx(data)
    if head[:5] == b"{\\rtf":
        from striprtf.striprtf import rtf_to_text
        return rtf_to_text(_extract_txt(data))
    if head[:4] == _OLE_MAGIC:
        docx = _soffice_to_docx(data)
        if docx is not None:
            return _extract_docx(docx)
    raise _ParseRejected(
        "Legacy .doc format needs conversion — save it as .docx and re-upload.")


def extract_text(filename: str, data: bytes) -> tuple[str, bool, str]:
    """Extract UTF-8 text from `data` based on `filename`'s extension.

    Returns (text, ok, error). Unsupported extension, an encrypted/corrupt file, or an
    extraction that yields no text (e.g. an image-only PDF) -> ok=False with a short error.
    Never raises: any low-level extractor error is caught and normalized to the error string
    so one bad file can't crash an upload of many.
    """
    ext = _ext(filename)
    if ext not in SUPPORTED_EXTS:
        return "", False, f"unsupported file type '{ext or filename}'"

    try:
        if ext in (".txt", ".md"):
            text = _extract_txt(data)
        elif ext == ".pdf":
            text = _extract_pdf(data)
        elif ext == ".docx":
            text = _extract_docx(data)
        elif ext == ".doc":
            text = _extract_doc(data)
        elif ext == ".xlsx":
            text = _extract_xlsx(data)
        elif ext == ".xls":
            text = _extract_xls(data)
        elif ext in (".csv", ".tsv"):
            text = _extract_csv(data, ext)
        else:  # .html / .htm
            text = _extract_html(data)
    except Exception as exc:  # noqa: BLE001 — normalize every extractor failure to the contract
        logger.info("parse failed for %s: %s", filename, exc)
        return "", False, f"could not extract text ({type(exc).__name__}): {str(exc)[:120]}"

    # Binary formats (PDF/DOCX/HTML) that extract to nothing are a FAILURE — an image-only
    # scan or a broken file the extractor opened but found no text in. Plain .txt/.md are
    # allowed to be legitimately empty (an empty note is still a valid, if useless, document).
    if not text.strip() and ext not in (".txt", ".md"):
        return "", False, "no extractable text (document may be image-only or empty)"
    return text, True, ""
