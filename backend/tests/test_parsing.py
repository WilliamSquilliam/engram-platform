"""Document parsing pipeline (app.parsing) + its wiring into upload/onboard, and the
config-gated GET /connectors contract. Pure-python extractors only (no torch/OCR)."""
import io
import zipfile

import pytest
from app import config
from app.parsing import SUPPORTED_EXTS, extract_text


# --- unit: the extractor per file type ------------------------------------------------

def test_txt_and_md_passthrough():
    text, ok, err = extract_text("note.txt", b"hello world")
    assert ok and err == "" and text == "hello world"
    text, ok, err = extract_text("readme.md", "# Title\n\nbody".encode("utf-8"))
    assert ok and text == "# Title\n\nbody"


def test_txt_decodes_non_utf8():
    # A cp1252 (Windows) file must not fail the upload — we fall through encodings.
    data = "café — résumé".encode("cp1252")
    text, ok, err = extract_text("legacy.txt", data)
    assert ok and "caf" in text and "sum" in text


def test_pdf_extracts_text():
    pytest.importorskip("pypdf")
    # A real one-page PDF with a text-showing operator (generated below without reportlab),
    # so pypdf's extract_text() returns the literal words.
    pdf_bytes = _tiny_text_pdf("Hello PDF world")
    text, ok, err = extract_text("doc.pdf", pdf_bytes)
    assert ok, err
    assert "Hello PDF" in text


def test_docx_extracts_text():
    pytest.importorskip("docx")
    from docx import Document as DocxDocument

    buf = io.BytesIO()
    d = DocxDocument()
    d.add_paragraph("First paragraph.")
    d.add_paragraph("Second paragraph.")
    table = d.add_table(rows=1, cols=1)
    table.rows[0].cells[0].text = "Cell content."
    d.save(buf)

    text, ok, err = extract_text("doc.docx", buf.getvalue())
    assert ok, err
    assert "First paragraph." in text
    assert "Second paragraph." in text
    assert "Cell content." in text  # table text is pulled too


def test_html_strips_tags():
    html = b"<html><head><style>.x{}</style></head><body><h1>Hi</h1>" \
           b"<script>ignore()</script><p>Body text</p></body></html>"
    text, ok, err = extract_text("page.html", html)
    assert ok, err
    assert "Hi" in text and "Body text" in text
    assert "ignore()" not in text and ".x{}" not in text  # script/style dropped


def test_unsupported_extension_fails():
    text, ok, err = extract_text("archive.zip", b"PK\x03\x04binary")
    assert not ok and text == ""
    assert "unsupported" in err.lower()


def test_corrupt_pdf_fails_gracefully():
    text, ok, err = extract_text("broken.pdf", b"%PDF-1.4 not really a pdf")
    assert not ok and text == "" and err  # short error, no exception


def test_image_only_pdf_is_failure():
    # A syntactically valid PDF with no extractable text -> failure (image-only signal).
    empty_pdf = _blank_pdf()
    text, ok, err = extract_text("scan.pdf", empty_pdf)
    assert not ok and "no extractable text" in err.lower()


# --- F5: decompression-bomb guards (clean parse failure, never an exception) -----------

def _docx_with_huge_declared_member() -> bytes:
    """A tiny .docx-shaped zip whose central directory DECLARES one member far over the 200 MB
    uncompressed ceiling. Only ~a few KB on disk (the classic zip-bomb shape). We forge the zip
    entry so file_size (the declared uncompressed size) is enormous without actually writing it."""
    import struct
    import zlib

    name = b"word/document.xml"
    payload = b"x" * 64
    comp = zlib.compress(payload)
    crc = zlib.crc32(payload) & 0xFFFFFFFF
    huge = 300 * 1024 * 1024  # 300 MB > the 200 MB guard

    # Local file header (declares the huge uncompressed size in the header too).
    lfh = struct.pack("<IHHHHHIIIHH", 0x04034B50, 20, 0, 8, 0, 0, crc,
                      len(comp), huge, len(name), 0) + name + comp
    # Central directory record mirroring it.
    cd = struct.pack("<IHHHHHHIIIHHHHHII", 0x02014B50, 20, 20, 0, 8, 0, 0, crc,
                     len(comp), huge, len(name), 0, 0, 0, 0, 0, 0) + name
    eocd = struct.pack("<IHHHHIIH", 0x06054B50, 0, 0, 1, 1, len(cd), len(lfh), 0)
    return lfh + cd + eocd


def test_docx_zip_bomb_rejected_as_parse_failure():
    """A DOCX declaring a huge uncompressed size is refused as a clean ok=False parse failure —
    never raises, never inflates."""
    bomb = _docx_with_huge_declared_member()
    text, ok, err = extract_text("bomb.docx", bomb)
    assert not ok and text == "" and err  # short reason, no exception


def test_docx_too_many_members_rejected():
    """A DOCX zip with an absurd member count is refused (member-count guard), as a parse failure."""
    from app import parsing

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        for i in range(parsing._MAX_DOCX_MEMBERS + 5):
            zf.writestr(f"part_{i}.xml", b"x")
    text, ok, err = extract_text("many.docx", buf.getvalue())
    assert not ok and text == "" and err


def test_pdf_over_page_limit_rejected(monkeypatch):
    """A PDF over the page ceiling is refused as a parse failure (page-count guard). We lower the
    ceiling instead of building a 2000-page PDF."""
    from app import parsing
    monkeypatch.setattr(parsing, "_MAX_PDF_PAGES", 1)

    from pypdf import PdfWriter
    w = PdfWriter()
    w.add_blank_page(width=100, height=100)
    w.add_blank_page(width=100, height=100)  # 2 pages > the patched limit of 1
    out = io.BytesIO()
    w.write(out)

    text, ok, err = extract_text("big.pdf", out.getvalue())
    assert not ok and text == "" and err


def test_supported_exts_matches_contract():
    assert SUPPORTED_EXTS == frozenset(
        {".txt", ".md", ".pdf", ".docx", ".html", ".htm"}
    )


# --- integration: parsing wired into upload + onboard ---------------------------------

def test_upload_sets_parse_status_and_persists_extracted_text(
    client, auth, make_corpus, mock_ml
):
    """Uploading a DOCX marks it parsed and persists EXTRACTED text so the onboard path
    reads words, not raw docx bytes."""
    from docx import Document as DocxDocument

    headers, _ = auth
    cid = make_corpus(headers)

    buf = io.BytesIO()
    d = DocxDocument()
    d.add_paragraph("The onboard path must see this sentence.")
    d.save(buf)

    docx_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    r = client.post(
        f"/corpora/{cid}/documents",
        files=[("files", ("handbook.docx", buf.getvalue(), docx_mime))],
        headers=headers,
    )
    assert r.status_code == 200, r.text
    doc = r.json()[0]
    assert doc["parse_status"] == "parsed"
    assert doc["parse_error"] is None

    # The onboard/train + retrieval path reads via storage.read_text — assert it now returns
    # the EXTRACTED words, not raw docx binary (which would start with "PK").
    from app.storage import storage
    text = storage.read_text(cid, "handbook.docx")
    assert "The onboard path must see this sentence." in text
    assert "PK" not in text[:4]


def test_upload_unparseable_marks_failed_with_error(client, auth, make_corpus):
    """A corrupt PDF is stored but marked failed with a short parse_error."""
    headers, _ = auth
    cid = make_corpus(headers)
    r = client.post(
        f"/corpora/{cid}/documents",
        files=[("files", ("broken.pdf", b"%PDF-1.4 garbage", "application/pdf"))],
        headers=headers,
    )
    assert r.status_code == 200, r.text
    doc = r.json()[0]
    assert doc["parse_status"] == "failed"
    assert doc["parse_error"]  # non-empty reason

    # A failed parse must NOT feed raw binary to onboarding: read_text returns "" (empty
    # sidecar), which the onboard path's non-empty filter drops.
    from app.storage import storage
    assert storage.read_text(cid, "broken.pdf") == ""


def test_upload_rejects_unsupported_type(client, auth, make_corpus):
    headers, _ = auth
    cid = make_corpus(headers)
    r = client.post(
        f"/corpora/{cid}/documents",
        files=[("files", ("data.zip", b"PK\x03\x04", "application/zip"))],
        headers=headers,
    )
    assert r.status_code == 400
    assert "unsupported" in r.text.lower()


def test_onboard_path_receives_extracted_text(client, auth, make_corpus, mock_ml):
    """End-to-end: the jobs onboard path builds its {doc_id, text} from storage.read_text,
    which returns EXTRACTED text for a parsed HTML doc (tags stripped)."""
    headers, _ = auth
    cid = make_corpus(headers)
    html = b"<html><body><h1>Policy</h1><p>Refunds within 30 days.</p>" \
           b"<script>x()</script></body></html>"
    r = client.post(
        f"/corpora/{cid}/documents",
        files=[("files", ("policy.html", html, "text/html"))],
        headers=headers,
    )
    assert r.status_code == 200, r.text
    assert r.json()[0]["parse_status"] == "parsed"

    # Reproduce exactly what routers/jobs._run_training builds for the ML client.
    from app.retrieval import cart_id_for
    from app.storage import storage
    filenames = storage.list_doc_filenames(cid)
    docs = [{"doc_id": cart_id_for("t", fn), "text": storage.read_text(cid, fn)}
            for fn in filenames]
    body = docs[0]["text"]
    assert "Refunds within 30 days." in body
    assert "<p>" not in body and "x()" not in body  # tags + script gone


# --- GET /connectors gating -----------------------------------------------------------

def test_connectors_requires_auth(client):
    assert client.get("/connectors").status_code == 401


def test_connectors_shape_and_default_gating(client, auth):
    """Contract: filesystem always available; google_drive + sharepoint gated off by default
    (conftest clears Google creds, so GDRIVE_ENABLED is False)."""
    headers, _ = auth
    r = client.get("/connectors", headers=headers)
    assert r.status_code == 200
    body = r.json()
    assert set(body) == {"connectors"}
    by_id = {c["id"]: c for c in body["connectors"]}
    assert set(by_id) == {"filesystem", "google_drive", "sharepoint"}
    for c in body["connectors"]:
        assert set(c) == {"id", "label", "available", "description"}
        assert isinstance(c["label"], str) and isinstance(c["description"], str)
    assert by_id["filesystem"]["available"] is True
    assert by_id["google_drive"]["available"] is False
    assert by_id["sharepoint"]["available"] is False


def test_connector_flips_available_when_creds_configured(client, auth, monkeypatch):
    """Setting OAuth creds (here via the config flag the registry reads fresh each call) flips
    a connector to available — the config-driven seam, no code change."""
    monkeypatch.setattr(config, "GDRIVE_ENABLED", True)
    headers, _ = auth
    body = client.get("/connectors", headers=headers).json()
    by_id = {c["id"]: c for c in body["connectors"]}
    assert by_id["google_drive"]["available"] is True
    assert by_id["sharepoint"]["available"] is False  # still off


# --- tiny PDF generators (no reportlab; a couple of literal PDFs) ----------------------

def _blank_pdf() -> bytes:
    """A minimal valid single-page PDF with no text content (pypdf reads it, extracts "")."""
    from pypdf import PdfWriter

    w = PdfWriter()
    w.add_blank_page(width=200, height=200)
    buf = io.BytesIO()
    w.write(buf)
    return buf.getvalue()


def _tiny_text_pdf(message: str) -> bytes:
    """Hand-rolled one-page PDF containing `message` as a text-showing operator, so pypdf's
    extract_text() returns it. Avoids a reportlab dependency for the test."""
    content = f"BT /F1 24 Tf 50 100 Td ({message}) Tj ET".encode("latin-1")
    objs = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 300 200] "
        b"/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>",
        b"<< /Length " + str(len(content)).encode() + b" >>\nstream\n" + content + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    out = bytearray(b"%PDF-1.4\n")
    offsets = []
    for i, body in enumerate(objs, start=1):
        offsets.append(len(out))
        out += f"{i} 0 obj\n".encode() + body + b"\nendobj\n"
    xref_pos = len(out)
    out += f"xref\n0 {len(objs) + 1}\n".encode()
    out += b"0000000000 65535 f \n"
    for off in offsets:
        out += f"{off:010d} 00000 n \n".encode()
    out += (f"trailer\n<< /Size {len(objs) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_pos}\n%%EOF").encode()
    return bytes(out)
